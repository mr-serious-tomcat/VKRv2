# Импорт необходимых библиотек
import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

import pandas as pd
from docx import Document


# Функция для парсинга диапазона строк
def parse_row_range(row_range_str):
    row_numbers = set()
    parts = row_range_str.split(',')
    for part in parts:
        if '-' in part:
            start, end = part.split('-')
            row_numbers.update(range(int(start), int(end)+1))
        else:
            row_numbers.add(int(part))
    return sorted(row_numbers)

# Функция для определения используемого шаблона
def determine_template(cell_content, template_dir):
    first_word = cell_content.strip().split(' ')[0].lower()  # Уже в нижнем регистре
    has_construction = re.search(r'\(в ред\..*?\)$', cell_content.strip(), re.IGNORECASE) is not None
    template_file_name = ""

    if first_word.startswith("постановление"):
        template_file_name = "ЭЗ постановление ШИ.docx" if has_construction else "ЭЗ постановление ШО.docx"
    elif first_word.startswith("приказ"):
        template_file_name = "ЭЗ приказ ШИ.docx" if has_construction else "ЭЗ приказ ШО.docx"
    elif first_word.startswith("закон"):
        template_file_name = "ЭЗ закон ШИ.docx" if has_construction else "ЭЗ закон ШО.docx"
    elif first_word.startswith("распоряжение"):
        template_file_name = "ЭЗ распоряжение ШИ.docx" if has_construction else "ЭЗ распоряжение ШО.docx"

    # Ищем шаблон без учета регистра
    for file in os.listdir(template_dir):
        if file.lower() == template_file_name.lower():
            return os.path.join(template_dir, file)
    return None

# Функция для замены конструкций в документе
def replace_placeholders(doc, replacements):
    def process_paragraph(paragraph, replacements):
        runs = paragraph.runs
        if not runs:
            return

        # Собираем полный текст и позиции
        full_text = []
        run_positions = []
        current_pos = 0
        for run in runs:
            text = run.text
            full_text.append(text)
            run_positions.append((current_pos, current_pos + len(text)))
            current_pos += len(text)

        full_text = ''.join(full_text)
        modified = False

        for key, value in replacements.items():
            if key not in full_text:
                continue

            # Поиск ключей
            start_idx = 0
            while True:
                start_idx = full_text.find(key, start_idx)
                if start_idx == -1:
                    break
                end_idx = start_idx + len(key)

                # Определяем, какие позиции затронуты
                affected_runs = []
                for i, (s, e) in enumerate(run_positions):
                    if s <= start_idx < e or s < end_idx <= e or (start_idx < s and end_idx > e):
                        affected_runs.append(i)

                if not affected_runs:
                    start_idx += 1
                    continue

                # Замена текста с сохранением стиля первой позиции
                first_run_idx = affected_runs[0]
                first_run = runs[first_run_idx]
                text_before = full_text[:start_idx]
                text_after = full_text[end_idx:]
                new_text = text_before + value + text_after

                # Очистить все позиции и распределить новый текст
                for run in runs:
                    run.text = ''

                # Сохраняем стиль первой позиции и вставляем новый текст
                first_run.text = new_text

                # Обновить данные для следующей итерации
                full_text = new_text
                modified = True
                break

        if modified:
            # Удаляем пустые позиции (если остались)
            for run in runs[::-1]:
                if not run.text:
                    p = run._element
                    p.getparent().remove(p)

    # Обработка всех параграфов
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, replacements)

    # Обработка заголовков, футеров и таблиц (аналогично)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, replacements)

# Функция для генерации имени документа
def generate_document_name(first_word_input, cell_content, template_path):
    first_word_input = first_word_input.strip()

    if 'постановление' in template_path.lower():
        template_type = 'постановление'
    elif 'приказ' in template_path.lower():
        template_type = 'приказ'
    elif 'закон' in template_path.lower():
        template_type = 'Закон'
    elif 'распоряжение' in template_path.lower():
        template_type = 'распоряжение'
    else:
        template_type = ''

    # Находим первую дату в формате дд.мм.гггг
    date_found = re.search(r'\b\d{2}\.\d{2}\.\d{4}\b', cell_content)
    date_found = date_found.group(0) if date_found else ''

    # Находим первую конструкцию '№...'
    num_found = re.search(r'№\s*([^\s)(]+)', cell_content)
    num_found = num_found.group(0) if num_found else ''

    additional_construction = ''
    if 'ШИ' in os.path.basename(template_path):
        additional = re.search(r'\(в ред\..*?\)', cell_content)
        additional_construction = additional.group(0) if additional else ''

    doc_name = f"{first_word_input} {template_type} {date_found} {num_found} {additional_construction}"
    # Удаляем лишние пробелы
    doc_name = ' '.join(doc_name.split())
    return doc_name

# Функция для замены запрещенных символов в имени файла
def replace_forbidden_characters(file_name):
    forbidden_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    for char in forbidden_chars:
        file_name = file_name.replace(char, '-')
    return file_name

# Класс приложения
class DocumentGeneratorApp:
    def __init__(self, master):
        self.sheet_combo = None
        self.master = master
        master.title("Генератор документов")
        master.geometry("530x400")

        # Переменные
        self.excel_file_path = tk.StringVar()
        self.sheet_name = tk.StringVar()
        self.row_range = tk.StringVar()
        self.template_dir = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.first_word_input = tk.StringVar()
        self.available_sheets = []

        # Создаем виджеты
        self.create_widgets()

    def create_widgets(self):
        # Файл Excel
        tk.Label(self.master, text="Шаг 1: Выберите файл Excel").grid(row=0, column=0, sticky='w', padx=10, pady=5)
        tk.Entry(self.master, textvariable=self.excel_file_path, width=50).grid(row=1, column=0, padx=10)
        tk.Button(self.master, text="Обзор", command=self.browse_excel_file).grid(row=1, column=1, padx=10)

        # Лист
        tk.Label(self.master, text="Шаг 2: Выберите лист из файла").grid(row=2, column=0, sticky='w', padx=10, pady=5)
        self.sheet_combo = ttk.Combobox(self.master, textvariable=self.sheet_name, values=self.available_sheets, state="readonly")
        self.sheet_combo.grid(row=3, column=0, padx=10)

        # Диапазон строк
        tk.Label(self.master, text="Шаг 3: Укажите строки (не из столбца А!) (например, 6-11 или 4,6-11)").grid(row=4, column=0, sticky='w', padx=10, pady=5)
        tk.Entry(self.master, textvariable=self.row_range, width=50).grid(row=5, column=0, padx=10)

        # Директория шаблонов
        tk.Label(self.master, text="Шаг 4: Выберите директорию шаблонов").grid(row=6, column=0, sticky='w', padx=10, pady=5)
        tk.Entry(self.master, textvariable=self.template_dir, width=50).grid(row=7, column=0, padx=10)
        tk.Button(self.master, text="Обзор", command=self.browse_template_dir).grid(row=7, column=1, padx=10)

        # Директория для сохранения документов
        tk.Label(self.master, text="Шаг 5: Выберите директорию для сохранения документов").grid(row=8, column=0, sticky='w', padx=10, pady=5)
        tk.Entry(self.master, textvariable=self.output_dir, width=50).grid(row=9, column=0, padx=10)
        tk.Button(self.master, text="Обзор", command=self.browse_output_dir).grid(row=9, column=1, padx=10)

        # Первое слово для названия документов
        tk.Label(self.master, text="Шаг 6: Введите первое слово для названия документов").grid(row=10, column=0, sticky='w', padx=10, pady=5)
        tk.Entry(self.master, textvariable=self.first_word_input, width=50).grid(row=11, column=0, padx=10)

        # Кнопка запуска
        tk.Button(self.master, text="Запустить процесс", command=self.start_process).grid(row=12, column=0, padx=10, pady=20)

    def browse_excel_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите файл Excel",
            filetypes=[("Excel files", "*.xlsx *.xlsm *.ods")])
        if file_path:
            self.excel_file_path.set(file_path)
            # Обновляем список листов
            self.update_sheet_names()

    def update_sheet_names(self):
        try:
            file_path = self.excel_file_path.get()
            if file_path.endswith('.ods'):
                df = pd.read_excel(file_path, engine='odf', sheet_name=None)
            else:
                df = pd.read_excel(file_path, sheet_name=None)
            self.available_sheets = list(df.keys())
            self.sheet_combo['values'] = self.available_sheets
            if self.available_sheets:
                self.sheet_combo.current(0)
                self.sheet_name.set(self.available_sheets[0])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось прочитать файл Excel: {e}")

    def browse_template_dir(self):
        dir_path = filedialog.askdirectory(title="Выберите директорию шаблонов")
        if dir_path:
            self.template_dir.set(dir_path)

    def browse_output_dir(self):
        dir_path = filedialog.askdirectory(title="Выберите директорию для сохранения документов")
        if dir_path:
            self.output_dir.set(dir_path)

    def start_process(self):
        # Проверка всех полей (обработка ошибок)
        if not self.excel_file_path.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите файл Excel.")
            return
        if not self.sheet_name.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите лист из файла Excel.")
            return
        if not self.row_range.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, введите диапазон строк.")
            return
        if not self.template_dir.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите директорию шаблонов.")
            return
        if not self.output_dir.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите директорию для сохранения документов.")
            return
        if not self.first_word_input.get():
            messagebox.showwarning("Предупреждение", "Пожалуйста, введите первое слово для названия документов.")
            return

        # Подтверждение запуска
        if not messagebox.askyesno("Подтверждение", "Нажмите 'Да' для начала создания документов."):
            return

        # Основной процесс создания
        try:
            excel_file = self.excel_file_path.get()
            sheet_selection = self.sheet_name.get()
            row_range_str = self.row_range.get()
            template_dir = self.template_dir.get()
            output_dir = self.output_dir.get()
            first_word_input = self.first_word_input.get()

            # Парсинг диапазона строк
            row_numbers = parse_row_range(row_range_str)

            # Проверка существования шаблонов
            template_files = [f for f in os.listdir(template_dir) if f.endswith('.docx')]
            if not template_files:
                messagebox.showerror("Ошибка", "В выбранной директории нет шаблонов в формате .docx.")
                return

            # Чтение файла Excel без заголовков
            if excel_file.endswith('.ods'):
                df = pd.read_excel(excel_file, engine='odf', sheet_name=sheet_selection, header=None)
            else:
                df = pd.read_excel(excel_file, sheet_name=sheet_selection, header=None)


            # Обработка каждой строки
            for idx in row_numbers:
                try:
                    row = df.iloc[idx - 1]  # Adjust for zero-based index
                    cellA = str(row[0]) if not pd.isna(row[0]) else ''
                    cellB = str(row[1]) if not pd.isna(row[1]) else ''
                    cellC = str(row[2]) if not pd.isna(row[2]) else ''
                    # Пропускаем, если cellC пустая
                    if not cellC.strip():
                        continue

                    # Определяем шаблон
                    template_path = determine_template(cellC, template_dir)
                    if not template_path:
                        messagebox.showwarning("Предупреждение", f"Не найден подходящий шаблон для строки {idx}.")
                        continue

                    # Открываем шаблон
                    doc = Document(template_path)

                    # Заменяем конструкции
                    replacements = {
                        '7-13cellC': cellC,
                        '7-14cellB': cellB,
                        '7-15cellA': cellA
                    }
                    replace_placeholders(doc, replacements)

                    # Формируем название документа
                    doc_name = generate_document_name(first_word_input, cellC, template_path)
                    doc_name = replace_forbidden_characters(doc_name)

                    # Проверяем, существует ли файл
                    save_path = os.path.join(output_dir, doc_name + '.docx')
                    if os.path.exists(save_path):
                        doc_name = 'D' + doc_name
                        save_path = os.path.join(output_dir, doc_name + '.docx')

                    # Сохраняем документ
                    doc.save(save_path)

                except Exception as e:
                    messagebox.showerror("Ошибка", f"Ошибка при обработке строки {idx}: {e}")

            messagebox.showinfo("Информация", "Генерация документов завершена.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось завершить процесс: {e}")

# Запуск приложения
if __name__ == '__main__':
    root = tk.Tk()
    app = DocumentGeneratorApp(root)
    root.mainloop()